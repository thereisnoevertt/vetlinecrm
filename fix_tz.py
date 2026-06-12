# -*- coding: utf-8 -*-
"""Преобразование ТЗ в приложение диплома:
- согласование с дипломом (6 менеджеров, без Telegram, шаблон, ENUM, нейтральная платформа);
- оформление по методичке (поля, отступ, заголовки, подписи таблиц);
- перенумерация таблиц в стиль приложения «Таблица А.N» + обновление ссылок;
- вложенные «Приложение А/Б» -> разделы 11/12;
- шапка «ПРИЛОЖЕНИЕ А» + заголовок, удаление титула с подписями и своего Содержания."""

import re
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = r'C:\vetline\TZ_Vetline_VK_v2.docx'
OUT = r'C:\vetline\TZ_Vetline_VK_v3.docx'
APP = 'А'

d = docx.Document(SRC)
body = d.element.body

# ─────────────────────────────────────────────────────────────────────────
# helpers
# ─────────────────────────────────────────────────────────────────────────
def all_paragraphs():
    for el in body.iter(qn('w:p')):
        yield Paragraph(el, d)

def ptext(p):
    return ''.join(r.text or '' for r in p.runs)

def set_text(p, newtext):
    if not p.runs:
        p.add_run(newtext); return
    p.runs[0].text = newtext
    for r in p.runs[1:]:
        r.text = ''

def is_in_table(p):
    el = p._p
    par = el.getparent()
    while par is not None:
        if par.tag == qn('w:tc'):
            return True
        par = par.getparent()
    return False

# ─────────────────────────────────────────────────────────────────────────
# 1. Поля и стиль Normal
# ─────────────────────────────────────────────────────────────────────────
for s in d.sections:
    s.left_margin = Cm(3); s.right_margin = Cm(1)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)

st = d.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(14)
pf = st.paragraph_format
pf.line_spacing = 1.5
pf.first_line_indent = Cm(1.25)
pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rpr = st.element.get_or_add_rPr()
rf = rpr.find(qn('w:rFonts'))
if rf is None:
    from docx.oxml import OxmlElement
    rf = OxmlElement('w:rFonts'); rpr.append(rf)
for k in ('w:ascii', 'w:hAnsi', 'w:cs'):
    rf.set(qn(k), 'Times New Roman')

# ─────────────────────────────────────────────────────────────────────────
# 2. Текстовые замены (содержание + нейтрализация платформы + эмодзи)
# ─────────────────────────────────────────────────────────────────────────
REPL = [
    # --- число менеджеров: диплом = 6 ---
    ('Отдел продаж: до 10 менеджеров', 'Отдел продаж: 6 менеджеров'),
    ('до 10 менеджеров', '6 менеджеров'),
    ('10 менеджеров + 1 директор + 1 администратор',
     '6 менеджеров + 1 директор + 1 администратор'),
    ('Менеджеры отдела продаж || до 10 человек', None),  # таблица — обработаем отдельно
    ('до 10 человек', '6 человек'),
    ('Заполнить справочник пользователей (10 менеджеров',
     'Заполнить справочник пользователей (6 менеджеров'),
    # --- Telegram отсутствует в дипломе ---
    ('переход с Telegram-бота на VK-бот, ', ''),
    ('переход с Telegram-бота на VK-бот', 'интеграция с VK-ботом'),
    # --- шаблон уведомления как в дипломе ---
    ('Команда «Ветлайн ВК».»', 'Команда «Ветлайн».»'),
    # --- эмодзи ---
    ('«📩 Уведомить»', '«Уведомить»'),
    ('«📦 Заказ поставщику»', '«Заказ поставщику»'),
    ('📩 ', ''), ('📦 ', ''), ('📩', ''), ('📦', ''),
    # --- ENUM статусов уведомления как в дипломе ---
    ('PENDING, SENT, ERROR', 'PENDING, SENT, RETRY_SCHEDULED, FAILED'),
    ('(SENT/ERROR/PENDING)', '(PENDING/SENT/RETRY_SCHEDULED/FAILED)'),
    # --- нейтрализация платформы (без openCRX и без жёсткого «монолита») ---
    ('Spring Boot 3.2.5', 'серверная Java-платформа (OpenJDK 17)'),
    ('Spring Web MVC + Thymeleaf 3.1 (server-side rendering)',
     'веб-слой с серверным рендерингом HTML-страниц'),
    ('Spring Security 6 + BCrypt (cost=12)',
     'подсистема безопасности с хешированием паролей BCrypt (cost=12)'),
    ('Spring Data JPA + Hibernate 6', 'слой ORM на основе стандарта JPA'),
    ('Flyway 10.10 (каталог db/migration/)',
     'средство версионирования схемы базы данных'),
    ('VK Java SDK 1.0.14 (Bots Long Poll API)',
     'клиентская библиотека VK Bot API (Bots Long Poll API)'),
    ('базовые знания Spring Boot, Maven',
     'базовые знания серверной Java-платформы и системы сборки'),
    ('миграций V1, V2, V3', 'миграций схемы базы данных'),
    ('миграции V1, V2, V3', 'миграции схемы базы данных'),
    ('Spring Boot Actuator (эндпоинт /actuator/health, /actuator/metrics)',
     'встроенные средства мониторинга платформы (проверка состояния и метрики)'),
    ('Документация Spring Boot 3.2.5: https://docs.spring.io/spring-boot/docs/3.2.5/reference/html/.',
     'Документация используемой серверной Java-платформы (актуальная версия).'),
    ('Все права реализованы декларативно через аннотации Spring Security и фильтры в SecurityConfig.',
     'Все права реализованы декларативно средствами подсистемы безопасности на уровне контроллеров и фильтров запросов.'),
    # --- вложенные приложения -> разделы (ссылки) ---
    ('чек-листу Приложения Б', 'чек-листу раздела 12'),
    ('чек-листа Приложения Б', 'чек-листа раздела 12'),
    ('Приложения Б', 'раздела 12'),
    ('Приложением Б', 'разделом 12'),
    ('Приложение Б', 'раздел 12'),
    ('Приложения А', 'раздела 11'),
    ('Приложение А', 'раздел 11'),
    # --- диапазон таблиц (спецслучай) ---
    ('таблицах 4.6 — 4.10', 'таблицах А.12 — А.16'),
    ('таблицах 4.6 – 4.10', 'таблицах А.12 – А.16'),
]
REPL = [(a, b) for (a, b) in REPL if b is not None]

# заголовки вложенных приложений -> разделы 11/12 (точные)
HEAD_RENAME = {
    'Приложение А. Матрица трассировки требований':
        '11 Матрица трассировки требований',
    'Приложение Б. Перечень контрольных вопросов для приёмки':
        '12 Перечень контрольных вопросов для приёмки',
}

for p in all_paragraphs():
    t = ptext(p)
    if not t:
        continue
    if t.strip() in HEAD_RENAME:
        set_text(p, HEAD_RENAME[t.strip()])
        continue
    nt = t
    for a, b in REPL:
        if a in nt:
            nt = nt.replace(a, b)
    if nt != t:
        set_text(p, nt)

# таблица 3.2: строка «Менеджеры отдела продаж | до 10 человек» уже покрыта 'до 10 человек'

# ─────────────────────────────────────────────────────────────────────────
# 3. Перенумерация таблиц -> «Таблица А.N» + подписи по методичке
# ─────────────────────────────────────────────────────────────────────────
order = ['1.1','1.2','2.1','2.2','3.1','3.2','4.1','4.2','4.3','4.4','4.5',
         '4.6','4.7','4.8','4.9','4.10','4.11','4.12','4.13','4.14','4.15',
         '4.16','5.1','7.1','8.1','9.1','А.1','Б.1']
tmap = {old: f'{APP}.{i+1}' for i, old in enumerate(order)}

cap_re = re.compile(r'^Таблица\s+([0-9A-ZА-Я]+\.[0-9]+)\s*[—–-]\s*(.*)$')
for p in all_paragraphs():
    t = ptext(p).strip()
    m = cap_re.match(t)
    if m and m.group(1) in tmap:
        set_text(p, f'Таблица {tmap[m.group(1)]} – {m.group(2)}')
        # формат подписи: слева, без отступа, перед 18 / после 0, обычный 14
        fmt = p.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(18); fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5
        for r in p.runs:
            r.bold = False; r.italic = False; r.font.size = Pt(14)

# ссылки на таблицы в тексте (после подписей): «таблиц* X.Y»
ref_re = re.compile(r'(табл\w*\.?\s+)([0-9][0-9]?\.[0-9]+)', re.IGNORECASE)
def ref_sub(mm):
    num = mm.group(2)
    return mm.group(1) + tmap.get(num, num)
for p in all_paragraphs():
    t = ptext(p)
    if 'абл' not in t:
        continue
    nt = ref_re.sub(ref_sub, t)
    if nt != t:
        set_text(p, nt)

# ─────────────────────────────────────────────────────────────────────────
# 4. Форматирование заголовков (по методичке)
# ─────────────────────────────────────────────────────────────────────────
h1 = re.compile(r'^\d+\s+\S')
h2 = re.compile(r'^\d+\.\d+(\.\d+)?\s+\S')
first_h1_seen = False
for p in all_paragraphs():
    if is_in_table(p):
        continue
    t = ptext(p).strip()
    runs = [r for r in p.runs if (r.text or '').strip()]
    if not runs:
        continue
    bold = any(r.bold for r in runs)
    fmt = p.paragraph_format
    if h1.match(t) and bold:
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0); fmt.space_after = Pt(18)
        fmt.page_break_before = first_h1_seen  # первый — без разрыва
        first_h1_seen = True
        for r in p.runs:
            r.bold = True; r.font.size = Pt(14)
    elif h2.match(t) and bold:
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(18); fmt.space_after = Pt(18)
        for r in p.runs:
            r.bold = True; r.font.size = Pt(14)

# ─────────────────────────────────────────────────────────────────────────
# 5. Форматирование тела таблиц (10–12 пт, одинарный, шапка по центру)
# ─────────────────────────────────────────────────────────────────────────
for tbl in d.tables:
    trs = tbl._tbl.findall(qn('w:tr'))
    for ridx, tr in enumerate(trs):
        for tc in tr.findall(qn('w:tc')):
            for pel in tc.findall(qn('w:p')):
                par = Paragraph(pel, tbl)
                fmt = par.paragraph_format
                fmt.first_line_indent = Cm(0)
                fmt.line_spacing = 1.0
                fmt.space_before = Pt(0); fmt.space_after = Pt(0)
                if ridx == 0:
                    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in par.runs:
                    r.font.size = Pt(11)
                    if ridx == 0:
                        r.bold = True

# ─────────────────────────────────────────────────────────────────────────
# 6. Удаление титула/подписей/Содержания + шапка «ПРИЛОЖЕНИЕ А»
# ─────────────────────────────────────────────────────────────────────────
# найти первый заголовок «1 Общие сведения» (полужирный)
target = None
for p in d.paragraphs:
    if ptext(p).strip() == '1 Общие сведения' and any(r.bold for r in p.runs):
        target = p._p; break
if target is not None:
    for p in list(d.paragraphs):
        if p._p is target:
            break
        p._p.getparent().remove(p._p)

# вставить шапку приложения и заголовок
def new_par_at_top(text, *, before, after, pagebreak):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
    r.font.name = 'Times New Roman'
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(before); fmt.space_after = Pt(after)
    if pagebreak:
        fmt.page_break_before = True
    body.remove(p._p); body.insert(0, p._p)
    return p

# порядок вставки: сначала заголовок (окажется 2-м), затем шапка (1-я)
new_par_at_top('Техническое задание на создание информационной системы '
               'взаимодействия с клиентами предприятия по продаже '
               'ветеринарного оборудования «Ветлайн ВК»',
               before=0, after=18, pagebreak=False)
new_par_at_top('ПРИЛОЖЕНИЕ ' + APP, before=0, after=18, pagebreak=True)

d.save(OUT)
print('Saved:', OUT)
